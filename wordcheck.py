from docx import Document
from collections import Counter
from tkinter import filedialog
import os

def compare_word_documents(doc1_path, doc2_path, max_repeats=20):
    # 读取两个Word文档
    doc1 = Document(doc1_path)
    doc2 = Document(doc2_path)
    checkword = Document()
    
    # 提取文档中的所有文本
    text1 = ''.join([p.text for p in doc1.paragraphs if p.text.strip()])
    text2 = ''.join([p.text for p in doc2.paragraphs if p.text.strip()])

    len1=len(text1)
    len2=len(text2)
    k=50

    for i in range(0, len1-k, 5):
        for j in range(len2):
            if text1[i:i+k]==text2[j:j+k]:
                print("文档1字符位",i,"与文档2字符位",j,"重复")
                print(text1[i:i+k])
                checkword.add_paragraph(f"文档1字符位{i}与文档2字符位{j}重复")
                checkword.add_paragraph(text1[i:i+k])
                checkword.add_paragraph()
    script_dir = os.path.dirname(doc1_path)
    os.chdir(script_dir)
    checkword.save('查重结果.docx')
 

# 替换为你的Word文档路径
doc1_path = filedialog.askopenfilename()
doc2_path = filedialog.askopenfilename()
#doc1_path = 'c:\\Users\\admin\\Desktop\\python\\word\\word1.docx'
#doc2_path = 'c:\\Users\\admin\\Desktop\\python\\word\\word2.docx'

# 调用函数比较文档
compare_word_documents(doc1_path, doc2_path)

# pyinstaller --noconsole --add-data "source_data.xlsx;." --onefile wordcheck.py
# pyinstaller --onefile wordcheck.py
